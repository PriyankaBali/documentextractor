"""DOCX document parsing."""

import io
from dataclasses import dataclass, field

from docx import Document as DocxDocument
from docx.table import Table


@dataclass
class DocxContent:
    """Content extracted from a DOCX file."""
    
    paragraphs: list[str]
    tables: list[list[list[str]]]  # List of tables, each table is list of rows
    full_text: str = ""
    metadata: dict = field(default_factory=dict)


class DocxParser:
    """Parse DOCX documents and extract content."""
    
    def parse(self, document) -> DocxContent:
        """Parse DOCX and extract text and tables.
        
        Args:
            document: LoadedDocument containing DOCX bytes.
            
        Returns:
            DocxContent with extracted data.
        """
        # Load DOCX from bytes
        docx_file = io.BytesIO(document.content)
        doc = DocxDocument(docx_file)
        
        # Extract paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = self._extract_table(table)
            if table_data:
                tables.append(table_data)
        
        # Build full text
        full_text = "\n".join(paragraphs)
        
        # Add table content to full text
        for table in tables:
            for row in table:
                full_text += "\n" + " | ".join(row)
        
        # Extract metadata
        metadata = self._extract_metadata(doc)
        
        return DocxContent(
            paragraphs=paragraphs,
            tables=tables,
            full_text=full_text,
            metadata=metadata,
        )
    
    def _extract_table(self, table: Table) -> list[list[str]]:
        """Extract data from a DOCX table.
        
        Args:
            table: python-docx Table object.
            
        Returns:
            List of rows, each row is a list of cell values.
        """
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Get cell text, handling merged cells
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        
        return table_data
    
    def _extract_metadata(self, doc: DocxDocument) -> dict:
        """Extract document metadata.
        
        Args:
            doc: python-docx Document object.
            
        Returns:
            Dictionary of metadata.
        """
        core_props = doc.core_properties
        
        return {
            "author": core_props.author or "",
            "title": core_props.title or "",
            "subject": core_props.subject or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }
